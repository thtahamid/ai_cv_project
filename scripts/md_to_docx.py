"""Convert docs/final_technical_report.md to a properly formatted Word document.

Produces:
  - Title page (centered)
  - Abstract section
  - Auto-updating Table of Contents
  - Numbered headings/subheadings
  - Justified body text
  - Centered figures with numbered captions at the bottom
  - Cleanly formatted tables
"""

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(r"C:\Users\mitah\github_projects\ai_cv_project")
SRC_MD = ROOT / "docs" / "final_technical_report.md"
DOCS_DIR = ROOT / "docs"
OUT_DOCX = DOCS_DIR / "final_technical_report.docx"


# ---------------------------- Short figure captions ----------------------------
# Maps the markdown image alt-text to a short, appropriate caption.
CAPTION_MAP = {
    "v1 class dist": "Class distribution — v1 dataset (Batch 04).",
    "v1 box area": "Box-area histogram — v1 dataset (Batch 04).",
    "v1 dims": "Image-dimension scatter — v1 dataset (Batch 04).",
    "v1 labels": "Label distribution — v1 dataset (Batch 04).",
    "v2 class dist": "Class distribution — v2 dataset (Batches 05/06).",
    "v2 box area": "Box-area histogram — v2 dataset (Batches 05/06).",
    "v2 dims": "Image-dimension scatter — v2 dataset (Batches 05/06).",
    "v2 labels": "Label distribution — v2 dataset (Batches 05/06).",
    "Batch 04 training curves": "Training curves — Batch 04.",
    "Batch 05 training curves": "Training curves — Batch 05.",
    "Batch 06 training curves": "Training curves — Batch 06.",
    "Batch 04 PR": "Precision–recall curve — Batch 04.",
    "Batch 05 PR": "Precision–recall curve — Batch 05.",
    "Batch 06 PR": "Precision–recall curve — Batch 06.",
    "Batch 04 F1": "F1 curve — Batch 04.",
    "Batch 05 F1": "F1 curve — Batch 05.",
    "Batch 06 F1": "F1 curve — Batch 06.",
    "Batch 04 CM": "Normalized confusion matrix — Batch 04.",
    "Batch 05 CM": "Normalized confusion matrix — Batch 05.",
    "Batch 06 CM": "Normalized confusion matrix — Batch 06.",
    "Batch 04 qualitative": "Qualitative predictions — Batch 04.",
    "Batch 05 qualitative": "Qualitative predictions — Batch 05.",
    "Batch 06 qualitative": "Qualitative predictions — Batch 06.",
}


# ---------------------------- Helpers ----------------------------

def set_cell_border(cell, **kwargs):
    """Set the borders on a table cell. kwargs keys: top, left, bottom, right."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            element = tcBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            element.set(qn("w:val"), kwargs[edge].get("val", "single"))
            element.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            element.set(qn("w:color"), kwargs[edge].get("color", "808080"))


def shade_cell(cell, color_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc):
    """Insert a Word field that renders an auto-updating Table of Contents."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and select 'Update Field' to refresh the table of contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`([^`]+)`")
INLINE_ITALIC = re.compile(r"(?<!\*)\*([^*\n]+)\*(?!\*)")


def add_runs_with_inline(paragraph, text: str):
    """Add runs to a paragraph, honoring **bold**, *italic*, and `code` inline."""
    # Tokenize: find all matches with their positions
    tokens = []  # list of (start, end, kind, content)
    for m in INLINE_BOLD.finditer(text):
        tokens.append((m.start(), m.end(), "bold", m.group(1)))
    for m in INLINE_CODE.finditer(text):
        tokens.append((m.start(), m.end(), "code", m.group(1)))
    # avoid overlapping italic with bold (** vs *) — italic detection runs on residual text
    tokens.sort()
    # filter overlaps
    filtered = []
    last_end = -1
    for t in tokens:
        if t[0] >= last_end:
            filtered.append(t)
            last_end = t[1]
    tokens = filtered

    cursor = 0
    for start, end, kind, content in tokens:
        if cursor < start:
            segment = text[cursor:start]
            # apply italic detection on plain segment
            _add_italic_segments(paragraph, segment)
        run = paragraph.add_run(content)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        cursor = end
    if cursor < len(text):
        _add_italic_segments(paragraph, text[cursor:])


def _add_italic_segments(paragraph, text: str):
    cursor = 0
    for m in INLINE_ITALIC.finditer(text):
        if cursor < m.start():
            paragraph.add_run(text[cursor:m.start()])
        run = paragraph.add_run(m.group(1))
        run.italic = True
        cursor = m.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<src>[^)]+)\)")
TABLE_ROW_RE = re.compile(r"^\s*\|(.+)\|\s*$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def parse_table_row(line: str):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


# ---------------------------- Styles ----------------------------

def configure_styles(doc):
    styles = doc.styles
    # Normal: justified, Times New Roman 11pt
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for h, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12), ("Heading 4", 11)]:
        st = styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        st.paragraph_format.space_before = Pt(14 if h == "Heading 1" else 10)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


# ---------------------------- Title page & abstract ----------------------------

def add_title_page(doc, meta: dict, title: str):
    # vertical spacer
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    run.font.name = "Calibri"

    doc.add_paragraph()
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = sep.add_run("Project 2 — Final Technical Report")
    sep_run.font.size = Pt(14)
    sep_run.font.italic = True
    sep_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(6):
        doc.add_paragraph()

    for label_key, label in [
        ("Project", "Project"),
        ("Author", "Author"),
        ("Date", "Date"),
        ("Hardware", "Hardware"),
        ("Repository", "Repository"),
    ]:
        if label_key in meta:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(12)
            r2 = p.add_run(meta[label_key])
            r2.font.size = Pt(12)

    add_page_break(doc)


def add_abstract(doc, abstract_paragraphs: list[str]):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Abstract")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    r.font.name = "Calibri"
    doc.add_paragraph()
    for para in abstract_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.6)
        add_runs_with_inline(p, para)
    add_page_break(doc)


def add_toc_page(doc):
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("Table of Contents")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    r.font.name = "Calibri"
    doc.add_paragraph()
    add_toc(doc)
    add_page_break(doc)


# ---------------------------- Figure handling ----------------------------

class FigureCounter:
    def __init__(self):
        self.n = 0

    def next(self):
        self.n += 1
        return self.n


def caption_for(alt: str) -> str:
    if alt in CAPTION_MAP:
        return CAPTION_MAP[alt]
    # fallback: titleize alt
    return alt.strip().capitalize() + "."


def add_figure(doc, image_path: Path, alt: str, fig_counter: FigureCounter, width_in=5.5):
    if not image_path.exists():
        # placeholder line
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"[Missing figure: {image_path}]")
        r.italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    n = fig_counter.next()
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r1 = cap.add_run(f"Figure {n}. ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = cap.add_run(caption_for(alt))
    r2.font.size = Pt(10)
    r2.italic = True


# ---------------------------- Table handling ----------------------------

def looks_like_image_only_table(rows: list[list[str]]) -> bool:
    """A 'figure grid' table is one whose data rows are entirely image links."""
    if len(rows) < 2:
        return False
    data_rows = rows[1:]  # skip header
    for row in data_rows:
        for cell in row:
            cell = cell.strip()
            if cell == "":
                continue
            if not IMAGE_RE.fullmatch(cell):
                return False
    # require at least one image
    for row in data_rows:
        for cell in row:
            if IMAGE_RE.fullmatch(cell.strip()):
                return True
    return False


def emit_image_grid(doc, rows: list[list[str]], fig_counter: FigureCounter):
    """Render image-only tables as a sequence of centered figures, one per image."""
    headers = rows[0]
    # iterate row-by-row, cell-by-cell — produce figures in reading order
    data_rows = rows[1:]
    for row in data_rows:
        for col_idx, cell in enumerate(row):
            m = IMAGE_RE.fullmatch(cell.strip())
            if not m:
                continue
            alt = m.group("alt").strip()
            src = m.group("src").strip()
            # caption — combine header (e.g., "Batch 04") with alt if alt is short
            header_label = headers[col_idx].strip() if col_idx < len(headers) else ""
            # resolve image path relative to docs/
            image_path = (DOCS_DIR / src).resolve()
            # width depends on column count for grouped figures
            width = 5.5 if len(headers) <= 2 else 5.0
            add_figure(doc, image_path, alt, fig_counter, width_in=width)


def emit_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row[j] if j < len(row) else ""
            # clear default empty paragraph
            cell.text = ""
            p = cell.paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs_with_inline(p, text)
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                shade_cell(cell, "1F3A5F")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                # right-align numeric cells
                stripped = text.strip().lstrip("*").rstrip("*")
                is_num = bool(re.match(r"^[-+−]?[0-9]+(\.[0-9]+)?( pp| %)?\*?$", stripped))
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_num else WD_ALIGN_PARAGRAPH.LEFT
                add_runs_with_inline(p, text)
                for run in p.runs:
                    run.font.size = Pt(10)
    # small space after
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------------------------- Markdown parsing & emission ----------------------------

def parse_and_emit(doc, lines: list[str], fig_counter: FigureCounter):
    i = 0
    in_code = False
    code_buf: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # fenced code blocks
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # horizontal rule — skip
        if stripped == "---":
            i += 1
            continue

        # blank
        if stripped == "":
            i += 1
            continue

        # standalone image
        m_img = IMAGE_RE.fullmatch(stripped)
        if m_img:
            image_path = (DOCS_DIR / m_img.group("src")).resolve()
            add_figure(doc, image_path, m_img.group("alt"), fig_counter)
            i += 1
            continue

        # headings
        if stripped.startswith("#"):
            level = 0
            while level < len(stripped) and stripped[level] == "#":
                level += 1
            heading_text = stripped[level:].strip()
            if level == 1:
                # title — skip; handled by title page
                i += 1
                continue
            level = min(level, 4)
            p = doc.add_heading(level=level - 1)  # Heading 1..3
            add_runs_with_inline(p, heading_text)
            i += 1
            continue

        # tables
        if TABLE_ROW_RE.match(line):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and (TABLE_ROW_RE.match(lines[j]) or TABLE_SEP_RE.match(lines[j])):
                table_lines.append(lines[j])
                j += 1
            # parse
            rows: list[list[str]] = []
            for tl in table_lines:
                if TABLE_SEP_RE.match(tl):
                    continue
                rows.append(parse_table_row(tl))
            if looks_like_image_only_table(rows):
                emit_image_grid(doc, rows, fig_counter)
            else:
                emit_table(doc, rows)
            i = j
            continue

        # unordered list
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_with_inline(p, stripped[2:].strip())
            i += 1
            continue

        # ordered list
        m_ol = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m_ol:
            p = doc.add_paragraph(style="List Number")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_with_inline(p, m_ol.group(2))
            i += 1
            continue

        # default — regular paragraph (justify)
        # collect continuation lines (until blank line / structural element)
        buf = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip("\n")
            nxt_s = nxt.strip()
            if (
                nxt_s == ""
                or nxt_s.startswith("#")
                or nxt_s.startswith("```")
                or nxt_s == "---"
                or TABLE_ROW_RE.match(nxt)
                or nxt_s.startswith(("- ", "* "))
                or re.match(r"^\d+\.\s+", nxt_s)
                or IMAGE_RE.fullmatch(nxt_s)
            ):
                break
            buf.append(nxt_s)
            j += 1
        text = " ".join(buf)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs_with_inline(p, text)
        i = j


# ---------------------------- Front-matter extraction ----------------------------

def split_frontmatter(text: str):
    lines = text.splitlines()
    # find title (first '# ')
    title = None
    meta = {}
    idx = 0
    for k, line in enumerate(lines):
        if line.startswith("# "):
            title = line[2:].strip()
            idx = k + 1
            break
    # metadata lines: "**Key:** value" until first horizontal rule or blank-then-heading
    while idx < len(lines):
        line = lines[idx].strip()
        if line == "" or line == "---":
            idx += 1
            continue
        m = re.match(r"^\*\*([^:]+):\*\*\s+(.*)$", line)
        if m:
            key = m.group(1).strip()
            value = m.group(2).strip()
            # strip inline code backticks for cleanliness
            value = value.replace("`", "")
            meta[key] = value
            idx += 1
            continue
        break

    # Extract abstract — between "## Abstract" and the next "---" or "## "
    abstract_paras: list[str] = []
    abs_start = None
    for k in range(idx, len(lines)):
        if lines[k].strip() == "## Abstract":
            abs_start = k + 1
            break
    if abs_start is not None:
        buf: list[str] = []
        cur_para: list[str] = []
        for k in range(abs_start, len(lines)):
            s = lines[k].rstrip("\n")
            stripped = s.strip()
            if stripped == "---" or stripped.startswith("## "):
                break
            if stripped == "":
                if cur_para:
                    abstract_paras.append(" ".join(cur_para).strip())
                    cur_para = []
            else:
                cur_para.append(stripped)
        if cur_para:
            abstract_paras.append(" ".join(cur_para).strip())

    # Body starts after the abstract section
    body_start = 0
    for k in range(len(lines)):
        if lines[k].strip() == "## Abstract":
            # advance to next "## " or "---"
            kk = k + 1
            while kk < len(lines):
                s = lines[kk].strip()
                if s.startswith("## ") and s != "## Abstract":
                    body_start = kk
                    break
                kk += 1
            break

    return title, meta, abstract_paras, lines[body_start:]


# ---------------------------- main ----------------------------

def main():
    md_text = SRC_MD.read_text(encoding="utf-8")
    title, meta, abstract_paras, body_lines = split_frontmatter(md_text)

    doc = Document()
    configure_styles(doc)

    # page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_title_page(doc, meta, title or "Final Technical Report")
    add_abstract(doc, abstract_paras)
    add_toc_page(doc)

    fig_counter = FigureCounter()
    parse_and_emit(doc, body_lines, fig_counter)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    print(f"Figures embedded: {fig_counter.n}")


if __name__ == "__main__":
    main()
